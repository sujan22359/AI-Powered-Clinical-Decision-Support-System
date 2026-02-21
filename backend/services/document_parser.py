"""
Document Parser Service for Clinical Report Analyzer

This module provides document parsing capabilities for PDF and DOCX files,
with robust error handling and content validation.
"""

import io
import logging
from pathlib import Path
from typing import Optional, Union, BinaryIO

import pdfplumber
import PyPDF2
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from backend.utils.logger import setup_logger


class DocumentParsingError(Exception):
    """Custom exception for document parsing errors"""
    pass


class DocumentParser:
    """
    Document parser for extracting text from PDF and DOCX files.
    
    Supports:
    - PDF text extraction using pdfplumber with PyPDF2 fallback
    - DOCX text extraction using python-docx
    - Content validation and error handling
    """
    
    def __init__(self):
        self.logger = setup_logger(__name__)
        self.supported_extensions = {'.pdf', '.docx'}
    
    def parse_document(self, file_content: Union[bytes, BinaryIO], filename: str) -> str:
        """
        Parse document and extract text content.
        
        Args:
            file_content: Document content as bytes or file-like object
            filename: Original filename to determine document type
            
        Returns:
            Extracted text content as string
            
        Raises:
            DocumentParsingError: If parsing fails or content is invalid
        """
        if not filename:
            raise DocumentParsingError("Filename is required to determine document type")
        
        # Get file extension
        file_extension = Path(filename).suffix.lower()
        
        if file_extension not in self.supported_extensions:
            raise DocumentParsingError(
                f"Unsupported file format: {file_extension}. "
                f"Supported formats: {', '.join(self.supported_extensions)}"
            )
        
        # Convert bytes to file-like object if needed
        if isinstance(file_content, bytes):
            file_obj = io.BytesIO(file_content)
        else:
            file_obj = file_content
        
        try:
            if file_extension == '.pdf':
                text = self._parse_pdf(file_obj)
            elif file_extension == '.docx':
                text = self._parse_docx(file_obj)
            else:
                raise DocumentParsingError(f"Unsupported file extension: {file_extension}")
            
            # Validate extracted content
            validated_text = self._validate_content(text, filename)
            
            self.logger.info(f"Successfully parsed document: {filename}")
            return validated_text
            
        except DocumentParsingError:
            raise
        except Exception as e:
            self.logger.error(f"Unexpected error parsing {filename}: {str(e)}")
            raise DocumentParsingError(f"Failed to parse document: {str(e)}")
    
    def _parse_pdf(self, file_obj: BinaryIO) -> str:
        """
        Extract text from PDF using pdfplumber with PyPDF2 fallback.
        
        Args:
            file_obj: PDF file object
            
        Returns:
            Extracted text content
            
        Raises:
            DocumentParsingError: If PDF parsing fails
        """
        text_content = ""
        
        try:
            # Primary method: pdfplumber
            self.logger.debug("Attempting PDF parsing with pdfplumber")
            file_obj.seek(0)
            
            with pdfplumber.open(file_obj) as pdf:
                if not pdf.pages:
                    raise DocumentParsingError("PDF contains no pages")
                
                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text + "\n"
                        else:
                            self.logger.warning(f"No text found on page {page_num}")
                    except Exception as e:
                        self.logger.warning(f"Error extracting text from page {page_num}: {str(e)}")
                        continue
            
            # If pdfplumber extracted content, return it
            if text_content.strip():
                self.logger.debug("Successfully extracted text using pdfplumber")
                return text_content
            
            # Fallback method: PyPDF2
            self.logger.debug("Falling back to PyPDF2 for PDF parsing")
            file_obj.seek(0)
            
            pdf_reader = PyPDF2.PdfReader(file_obj)
            
            if not pdf_reader.pages:
                raise DocumentParsingError("PDF contains no pages")
            
            fallback_text = ""
            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        fallback_text += page_text + "\n"
                except Exception as e:
                    self.logger.warning(f"PyPDF2 error on page {page_num}: {str(e)}")
                    continue
            
            if fallback_text.strip():
                self.logger.debug("Successfully extracted text using PyPDF2 fallback")
                return fallback_text
            
            raise DocumentParsingError("No text content could be extracted from PDF")
            
        except DocumentParsingError:
            raise
        except Exception as e:
            self.logger.error(f"PDF parsing failed: {str(e)}")
            raise DocumentParsingError(f"Failed to parse PDF: {str(e)}")
    
    def _parse_docx(self, file_obj: BinaryIO) -> str:
        """
        Extract text from DOCX file using python-docx.
        
        Args:
            file_obj: DOCX file object
            
        Returns:
            Extracted text content
            
        Raises:
            DocumentParsingError: If DOCX parsing fails
        """
        try:
            self.logger.debug("Attempting DOCX parsing with python-docx")
            file_obj.seek(0)
            
            doc = Document(file_obj)
            
            text_content = ""
            paragraph_count = 0
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                paragraph_text = paragraph.text.strip()
                if paragraph_text:
                    text_content += paragraph_text + "\n"
                    paragraph_count += 1
            
            # Extract text from tables
            table_count = 0
            for table in doc.tables:
                table_count += 1
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_content += " | ".join(row_text) + "\n"
            
            self.logger.debug(f"Extracted text from {paragraph_count} paragraphs and {table_count} tables")
            
            if not text_content.strip():
                raise DocumentParsingError("No text content found in DOCX document")
            
            return text_content
            
        except PackageNotFoundError:
            raise DocumentParsingError("Invalid DOCX file format")
        except DocumentParsingError:
            raise
        except Exception as e:
            self.logger.error(f"DOCX parsing failed: {str(e)}")
            raise DocumentParsingError(f"Failed to parse DOCX: {str(e)}")
    
    def _validate_content(self, text: str, filename: str) -> str:
        """
        Validate extracted text content.
        
        Args:
            text: Extracted text content
            filename: Original filename for logging
            
        Returns:
            Validated and cleaned text content
            
        Raises:
            DocumentParsingError: If content validation fails
        """
        if not text or not text.strip():
            raise DocumentParsingError("Document appears to be empty or contains no readable text")
        
        # Clean up the text
        cleaned_text = text.strip()
        
        # Basic content validation
        if len(cleaned_text) < 10:
            raise DocumentParsingError("Document content is too short to be meaningful")
        
        # Check for reasonable character distribution (not just whitespace/special chars)
        alphanumeric_chars = sum(1 for c in cleaned_text if c.isalnum())
        if alphanumeric_chars < len(cleaned_text) * 0.1:  # At least 10% alphanumeric
            raise DocumentParsingError("Document content appears to be corrupted or unreadable")
        
        self.logger.debug(f"Content validation passed for {filename}: {len(cleaned_text)} characters")
        return cleaned_text
    
    def is_supported_format(self, filename: str) -> bool:
        """
        Check if file format is supported.
        
        Args:
            filename: Filename to check
            
        Returns:
            True if format is supported, False otherwise
        """
        if not filename:
            return False
        
        file_extension = Path(filename).suffix.lower()
        return file_extension in self.supported_extensions
    
    def get_supported_formats(self) -> set:
        """
        Get set of supported file formats.
        
        Returns:
            Set of supported file extensions
        """
        return self.supported_extensions.copy()